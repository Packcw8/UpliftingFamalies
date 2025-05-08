from docx import Document

def render_checkbox(selected: str, options: list[str]) -> str:
    return "    ".join([f"{{*}} {opt}" if opt == selected else f"{{ }} {opt}" for opt in options])

def fill_template(data):
    doc = Document("app/templates/visit_template.docx")


    replacements = {
        "{{case_name}}": data.case_name,
        "{{case_number}}": data.case_number,
        "{{service_date}}": data.service_date,
        "{{service_provided}}": data.service_provided,
        "{{code}}": ", ".join(data.codes),
        "{{skill_deficit}}": data.skill_deficit,
        "{{developed_skill}}": data.developed_skill,
        "{{participants}}": data.participants,
        "{{start_time}}": data.start_time,
        "{{stop_time}}": data.stop_time,
        "{{units}}": str(data.units),
        "{{Summary}}": data.summary,
        "{{clients_progress}}": data.clients_progress,
        "{{miles}}": str(data.miles),
        "{{at_units}}": str(data.at_units),
        "{{it_units}}": str(data.it_units),
        "{{signature}}": data.signature,
        "{{completion_status}}": render_checkbox(data.completion_status, ["Completed", "Attempted"]),
        "{{safety_checkbox}}": render_checkbox(data.safety_checkbox, ["Safety Concerns", "No Safety Concerns"]),
        "{{abuse_checkbox}}": render_checkbox(data.abuse_checkbox, ["Abuse or Neglect", "No Abuse or Neglect"]),
        "{{location_checkbox}}": render_checkbox(data.location_checkbox, ["Home", "DHHR", "Court", "Other"])
    }

    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)

    for i in range(1, 4):
        route = getattr(data, f"route_{i}", {})
        doc_text_replace(doc, f"{{{{from_{i}}}}}", route.get("from", ""))
        doc_text_replace(doc, f"{{{{to_{i}}}}}", route.get("to", ""))
        doc_text_replace(doc, f"{{{{at_start{i}}}}}", route.get("at_start", ""))
        doc_text_replace(doc, f"{{{{at_stop{i}}}}}", route.get("at_stop", ""))
        doc_text_replace(doc, f"{{{{it_start{i}}}}}", route.get("it_start", ""))
        doc_text_replace(doc, f"{{{{it_stop{i}}}}}", route.get("it_stop", ""))

    return doc

def doc_text_replace(doc: Document, placeholder: str, value: str):
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, value)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if placeholder in c.text:
                    c.text = c.text.replace(placeholder, value)

